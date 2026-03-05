# skills/export_skill.py
import os
from typing import Dict, Any, List

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

from docx import Document
from config import EXPORT_DIR


# --- Register a built-in CJK font (no external ttf needed) ---
# This solves "■■■" squares for Chinese text.
pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def _p(text: str) -> str:
    """Safe paragraph text."""
    return (text or "").replace("\n", "<br/>")


def _kv_table(rows: List[List[str]]) -> Table:
    t = Table(rows, colWidths=[90, 430])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.whitesmoke, colors.white]),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def export_case_pdf(case: Dict[str, Any]) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    out_path = os.path.join(EXPORT_DIR, f"{case['case_id']}.pdf")

    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "title_cn", parent=styles["Title"], fontName="STSong-Light", fontSize=16, leading=20
    )
    h_style = ParagraphStyle(
        "h_cn", parent=styles["Heading2"], fontName="STSong-Light", fontSize=12, leading=16, spaceBefore=8, spaceAfter=6
    )
    body_style = ParagraphStyle(
        "body_cn", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=15
    )
    small_style = ParagraphStyle(
        "small_cn", parent=styles["BodyText"], fontName="STSong-Light", fontSize=9.5, leading=13, textColor=colors.grey
    )

    story = []

    case_id = case.get("case_id", "")
    stage = case.get("stage", "")
    turns = case.get("turns", [])
    rounds = sum(1 for t in turns if t.get("role") == "user")

    # Header (match your reference style)
    story.append(Paragraph("中医问诊记录（自动导出）", title_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"case_id：{case_id} | 阶段：{stage} | 轮次：{rounds}", small_style))
    story.append(Spacer(1, 12))

    fd = case.get("four_diagnosis", {})
    inq = (fd.get("inquiry") or {})
    tongue = (fd.get("inspection") or {}).get("tongue", "") or inq.get("tongue", "")
    pulse = (fd.get("palpation") or {}).get("pulse", "") or inq.get("pulse", "")

    # 1) 主诉与病程
    story.append(Paragraph("一、主诉与病程", h_style))
    chief = inq.get("chief_complaint", "") or "—"
    present = inq.get("present_illness", "") or "—"
    story.append(_kv_table([
        ["主诉", chief],
        ["现病史", present],
    ]))
    story.append(Spacer(1, 10))

    # 2) 四诊信息
    story.append(Paragraph("二、四诊信息", h_style))
    symptoms = inq.get("symptoms", [])
    symptoms_txt = "、".join(symptoms) if isinstance(symptoms, list) and symptoms else (inq.get("symptoms") or "—")

    story.append(_kv_table([
        ["舌象", tongue or "—"],
        ["脉象", pulse or "—"],
        ["寒热", inq.get("cold_heat", "") or "—"],
        ["汗出", inq.get("sweat", "") or "—"],
        ["口渴", inq.get("thirst", "") or "—"],
        ["睡眠", inq.get("sleep", "") or "—"],
        ["纳食", inq.get("appetite", "") or "—"],
        ["大便", inq.get("stool", "") or "—"],
        ["小便", inq.get("urine", "") or "—"],
        ["疼痛", inq.get("pain", "") or "—"],
        ["伴随症状", symptoms_txt or "—"],
        ["其他", "—"],
    ]))
    story.append(Spacer(1, 10))

    # 3) 辨证假设（如果你有 diagnosis / patterns）
    story.append(Paragraph("三、辨证假设", h_style))
    diag = case.get("diagnosis", {}) or {}
    diag_text = diag.get("text", "") or "—"
    conf = diag.get("confidence", None)
    conf_str = f"{conf:.2f}" if isinstance(conf, (int, float)) else "—"
    patterns = diag.get("patterns", [])
    patterns_txt = "；".join(patterns) if isinstance(patterns, list) and patterns else "—"

    # 小表格：证型/方向 | 置信度 | 依据
    basis = ""
    # 兼容你不同版本的字段
    if isinstance(diag, dict):
        basis = diag.get("brief_basis", "") or diag.get("basis", "") or ""
    basis = basis or "—"

    t = Table(
        [["证型/方向", "依据（简述）", "置信度"],
         [patterns_txt, basis, conf_str]],
        colWidths=[160, 280, 80]
    )
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"诊断说明：{_p(diag_text)}", body_style))
    story.append(Spacer(1, 10))

    # 4) 治则与建议
    story.append(Paragraph("四、治则与建议", h_style))
    advice = case.get("advice", "") or "—"
    story.append(Paragraph(_p(advice), body_style))
    story.append(Spacer(1, 10))

    # 5) 风险提示
    story.append(Paragraph("五、风险提示", h_style))
    risk = case.get("risk", {}) or {}
    risk_level = risk.get("level", "none")
    risk_reasons = risk.get("reasons", [])
    rr = "、".join(risk_reasons) if isinstance(risk_reasons, list) and risk_reasons else "—"
    story.append(_kv_table([
        ["风险等级", str(risk_level)],
        ["警示/红旗信号", rr],
    ]))
    story.append(Spacer(1, 10))

    # 6) RAG 证据命中（可解释性）——你参考PDF里有这一节 :contentReference[oaicite:2]{index=2}
    story.append(Paragraph("六、RAG 证据命中（可解释性）", h_style))
    evidence = case.get("evidence", []) or []
    if evidence:
        for i, e in enumerate(evidence[:12], start=1):
            src = e.get("source", "") or "-"
            score = e.get("score", None)
            score_s = f"{score:.3f}" if isinstance(score, (int, float)) else "-"
            txt = (e.get("text", "") or "").strip()
            snippet = txt[:260].replace("\n", " ")
            story.append(Paragraph(f"[{i}] {src} score={score_s}", body_style))
            story.append(Paragraph(_p(snippet), small_style))
            story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("—", body_style))

    doc.build(story)
    return out_path


def export_case_docx(case: Dict[str, Any]) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    out_path = os.path.join(EXPORT_DIR, f"{case['case_id']}.docx")
    doc = Document()
    doc.add_heading("中医问诊记录（自动导出）", level=1)
    doc.add_paragraph(f"case_id：{case.get('case_id','')} | 阶段：{case.get('stage','')}")

    fd = case.get("four_diagnosis", {})
    inq = (fd.get("inquiry") or {})
    tongue = (fd.get("inspection") or {}).get("tongue", "") or inq.get("tongue", "")
    pulse = (fd.get("palpation") or {}).get("pulse", "") or inq.get("pulse", "")

    doc.add_heading("一、主诉与病程", level=2)
    doc.add_paragraph(f"主诉：{inq.get('chief_complaint','') or '—'}")
    doc.add_paragraph(f"现病史：{inq.get('present_illness','') or '—'}")

    doc.add_heading("二、四诊信息", level=2)
    doc.add_paragraph(f"舌象：{tongue or '—'}")
    doc.add_paragraph(f"脉象：{pulse or '—'}")
    for k in ["cold_heat","sweat","thirst","sleep","appetite","stool","urine","pain"]:
        doc.add_paragraph(f"{k}：{inq.get(k,'') or '—'}")
    syms = inq.get("symptoms", [])
    doc.add_paragraph(f"伴随症状：{'、'.join(syms) if isinstance(syms,list) and syms else '—'}")

    doc.add_heading("三、辨证假设", level=2)
    diag = case.get("diagnosis", {}) or {}
    doc.add_paragraph(f"证型/方向：{'；'.join(diag.get('patterns',[]) or []) or '—'}")
    doc.add_paragraph(f"置信度：{diag.get('confidence','—')}")
    doc.add_paragraph(f"结论：{diag.get('text','') or '—'}")

    doc.add_heading("四、治则与建议", level=2)
    doc.add_paragraph(case.get("advice","") or "—")

    doc.add_heading("五、风险提示", level=2)
    risk = case.get("risk", {}) or {}
    doc.add_paragraph(f"风险等级：{risk.get('level','none')}")
    doc.add_paragraph(f"原因：{risk.get('reasons',[]) or '—'}")

    doc.add_heading("六、RAG证据命中", level=2)
    ev = case.get("evidence", []) or []
    for i, e in enumerate(ev[:12], start=1):
        doc.add_paragraph(f"[{i}] {e.get('source','-')} score={e.get('score','-')}")
        doc.add_paragraph((e.get("text","") or "")[:400])

    doc.save(out_path)
    return out_path